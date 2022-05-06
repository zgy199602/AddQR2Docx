###########################################################################################
# 将1个人的健康码行程码等添加名字后放到一个word中，能处理1个人有1到4个码的情况
# 希望各种码按人名保存，保存的目录结构如下：
# --rootdir
#   --01xxx(人名1)                    人名有顺序需求的话，使用序号表明顺序
#      --01xxx.jpg(健康码截图1)       健康码有顺序需求的话，使用序号表明顺序
#      --02xxx.jpg(健康码截图2)       注意序号最好是01、02,11、12等，而不是1、2
#      --03xxx.jpg(健康码截图3)
#      --04xxx.jpg(健康码截图4)
#   --02xxx(人名2)
#      --01xxx.jpg(健康码截图1)
#      ......
#   ......
# 最后会在rootdir下生成一个“test.docx”word文件，检查后改名即可
# 同时会在每个人的名字下生成“NameAdd”文件夹，保存了添加名字后的健康码图片
###########################################################################################
import os
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Cm,Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from string import digits
from tkinter import filedialog

#rootdir = "C:\\users\\xxx\\Desktop\\健康码和行程码\\20220222"  #根据情况修改
rootdir = filedialog.askdirectory()

os.chdir(rootdir)
names = os.listdir()
names.sort()
#图片缩放和添加名字
for name in names:
    cwdir = os.path.join(rootdir,name)
    if os.path.isdir(cwdir):
        os.chdir(cwdir)
        if not os.path.exists("NameAdd"):
            os.makedirs("NameAdd")
        name1 = name.translate(str.maketrans('','',digits)) #去掉文件夹名字中的数字
        images = os.listdir()
        for image in images:
            if os.path.isdir(image):continue
            else:
                im = Image.open(image)
                im = im.resize((204,456),Image.Resampling.LANCZOS)
                TruetypeFonts = ImageFont.truetype("Fonts/simhei.ttf",20) #黑体，字号20

                draw = ImageDraw.Draw(im)
                length = 20 * len(name1) #根据名字字数来决定空白区域大小
                draw.rectangle((70,30,70+length,50),fill="white") #先画一个空白区域用来填充名字
                #name = (name[0] + " " + name[1]) if (len(name)==2) else name
                draw.text((70,30),name1,fill="black",font=TruetypeFonts)                 
                im.save(os.path.join("NameAdd",image),quality=96)

# 将生成的图片写入word
doc = docx.Document()
# 改变纸张大小为A4
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)

for i, name in enumerate(names):
    os.chdir(os.path.join(rootdir,name,"NameAdd"))
    pics = os.listdir()
    pics.sort()
    doc.add_paragraph()
    doc.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER #段落居中

    pics_num = len(pics)
    if(pics_num<2 or pics_num>=5):
        print("Error:该程序只能处理1、2、3、4张图片，目前为%d张，不符合要求" % pics_num)
        break
    if((pics_num>2) and (pics_num<5)):  #三或四张图每张大小的设置，让A4一页刚好能放下4张图片
        A4h = 11.94
        A4w = 5.37
    if(pics_num==2):                    #两张图每张大小的设置，让A4一页放2张图片
        A4h = 15.82
        A4w = 7.26    
    if(pics_num==1):                    #一张图每张大小的设置，让A4一页放1张图片
        A4h = 24
        A4w = 11   

    for pic in pics:
        doc.paragraphs[i].add_run().add_picture(pic,height=Cm(A4h),width=Cm(A4w))
    
os.chdir(rootdir)
doc.save("test.docx")